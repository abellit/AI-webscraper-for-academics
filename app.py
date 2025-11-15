import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from time import gmtime, strftime
from fpdf import FPDF
from docx import Document
from utils.scrapper import scrape_website, extract_main_content, clean_main_content, split_dom_content
from utils.parse import parse_with_ollama

# Streamlit app configuration
st.title("AI Web Scraper")

url_input_method = st.radio("Select URL Input Method:", ("Single URL", "Multiple URLs"))

parsed_results = ""

if url_input_method == "Single URL":
    input_field = st.text_input("Enter a Website URL: ")
    url = [input_field] if input_field else []
else:
    input_field = st.text_area("Enter a Website URL: ", height=100)
    urls = [url.strip() for url in input_field.split("\n") if url.strip()]

# # Scrape button fundtionality
# if st.button("Scrape Site"):
#     st.write("Scraping the website")
#     if url_input_method == "Multiple URLs":
#         all_dom_contents = []
#         for url in urls:
#             try:
#                 result = scrape_website(url)
#                 main_content = extract_main_content(result)
#                 cleaned_content = clean_main_content(main_content)
#                 all_dom_contents.append(f"Content from {url}:\n{cleaned_content}\n\n")
#                 st.session_state.dom_content = "\n".join(all_dom_contents)
#             except Exception as e:
#                 st.error(f"Error scraping {url}: {e}")
#     result = scrape_website(url)
#     main_content = extract_main_content(result)
#     cleaned_content = clean_main_content(main_content)
#     st.session_state.dom_content = cleaned_content

#     with st.expander("View DOM Content"):
#         st.text_area("DOM Content", cleaned_content, height=300)

# Scrape button functionality
if st.button("Scrape Site"):
    st.write("Scraping the website...")

    if url_input_method == "Single URL":
        # 'url' is the list from line 30, e.g., ['https://google.com'] or []
        if url: 
            single_url_to_scrape = url[0] # Get the FIRST item (the string)
            
            if single_url_to_scrape: # Check it's not an empty string
                try:
                    result = scrape_website(single_url_to_scrape) # Pass the STRING
                    main_content = extract_main_content(result)
                    cleaned_content = clean_main_content(main_content)
                    st.session_state.dom_content = cleaned_content
                    
                    with st.expander("View DOM Content"):
                        st.text_area("DOM Content", cleaned_content, height=300)
                except Exception as e:
                    st.error(f"Error scraping {single_url_to_scrape}: {e}")
            else:
                st.warning("Please enter a URL.")
        else:
            st.warning("Please enter a URL.")
            
    else: # This means url_input_method == "Multiple URLs"
        # 'urls' is the list from line 33
        if urls: 
            all_dom_contents = []
            for individual_url in urls: # Loop through the 'urls' (plural) list
                try:
                    result = scrape_website(individual_url) # Pass the string
                    main_content = extract_main_content(result)
                    cleaned_content = clean_main_content(main_content)
                    all_dom_contents.append(f"Content from {individual_url}:\n{cleaned_content}\n\n")
                    st.success(f"Successfully scraped {individual_url}")
                except Exception as e:
                    st.error(f"Error scraping {individual_url}: {e}")
            
            if all_dom_contents:
                st.session_state.dom_content = "\n".join(all_dom_contents)
                with st.expander("View DOM Content"):
                    st.text_area("DOM Content", st.session_state.dom_content, height=300)
        else:
            st.warning("Please enter at least one URL in the text area.")

# Parsing functionality  
if "dom_content" in st.session_state:
    parse_description = st.text_area("Describe what you want to parse.")
    
    if st.button("Parse Content"):
        if parse_description:
            st.write("Parsing content...")
            dom_chunks = split_dom_content(st.session_state.dom_content)
            results = parse_with_ollama(dom_chunks, parse_description)
            st.write(results)
            st.session_state.parsed_results = results
    if "parsed_results" in st.session_state:
        st.write("Parsing complete. Ready for download.")

st.write("\n")
st.write("To download the parsed results, enter the file name (including file extension)")


time = strftime("%a, %d %b %Y", gmtime())
file_name = st.text_input("Enter file name for download (e.g., results.csv):", )
valid_extensions = [".txt", ".csv", ".json", ".xlsx", ".pdf", ".docx", ".md"]

download_data = b""
file_mime = "application/octet-stream"
valid_for_download = False

validate_extension = False
if file_name and any(file_name.endswith(ext) for ext in valid_extensions):  
    validate_extension = True
else:
    st.error(f"Invalid file extension. Please use one of the following: {', '.join(valid_extensions)}")


if "parsed_results" in st.session_state and validate_extension:
    parsed_results = st.session_state.parsed_results


    try:
        if file_name.endswith(".csv"):
            df = pd.DataFrame([x.split(",") for x in parsed_results.split("\n")])
            download_data = df.to_csv(index=False).encode("utf-8")
            file_mime = "text/csv"
            valid_for_download = True

        elif file_name.endswith(".xlsx"):
            df = pd.DataFrame([x.split(",") for x in parsed_results.split("\n")])
            output = BytesIO()
            with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                df.to_excel(writer, index=False)
            download_data = output.getvalue()
            file_mime = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            valid_for_download = True

        elif file_name.endswith(".json"):
            df = pd.json_normalize([x.split(",") for x in parsed_results.split("\n")])
            download_data = df.to_json(orient="records").encode("utf-8")
            file_mime = "application/json"
            valid_for_download = True

        elif file_name.endswith(".txt") or file_name.endswith(".md"):
            download_data = parsed_results.encode("utf-8")
            file_mime = "text/plain"
            valid_for_download = True

        elif file_name.endswith(".pdf"):
            class PDF(FPDF):

                def header(self, time):
                    self.set_font("Arial", "B", 12)
                    self.cell(0, 10, time, 0, 1, "L")
                
                def main(self, text):
                    self.set_font("Arial", "", 12)
                    self.title("Parsed Results", ln=True, align="C", size=18, thickness=2)
                    self.multi_cell(0, 10, text)
                    self.ln()

                def footer(self):
                    self.set_y(-15)
                    self.set_font("Arial", "I", 8)
                    self.cell(0, 10, f"Page {self.page_no()}", 0, 0, "R")

            pdf = PDF()
            pdf.add_page()
            pdf.header(time)
            pdf.main(parsed_results)
            download_data = pdf.output(dest="S").encode("Latin-1")
            file_mime = "application/pdf"
            valid_for_download = True

        elif file_name.endswith(".docx"):
            
            document = Document()
            document.add_heading(time, 0)
            document.add_heading("Parsed Results", level=1)
            document.add_paragraph(parsed_results)

            doc_io = BytesIO()
            document.save(doc_io)
            download_data = doc_io.getvalue()
            file_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            valid_for_download = True

    except Exception as e:
        st.error(f"Error converting parsed results to {file_name} file: {e}")
        valid_for_download = False


st.download_button(
    label="Download Parsed Results",
    data=download_data,
    file_name=file_name if file_name else "parsed_results.txt",
    mime=file_mime,
    disabled=not valid_for_download
)

# Instructions and notes
# st.markdown(""
#     "### Instructions:\n"
#     "1. Enter a valid website URL in the input box.\n"
#     "2. Click the 'Scrape Site' button to fetch the website content.\n"
#     "3. Review the DOM content in the expandable section.\n"
#     "4. Provide a description of the information you want to extract in the text area.\n"
#     "5. Click the 'Parse Content' button to extract the specified information.\n"
#     "6. The parsed results will be displayed below the button.\n"
#     "7. If no information matches the description, an empty string will be returned.\n"
#     "8. Ensure the website is accessible and the URL is correct.\n"
#     "9. If you encounter any issues, check the console for error messages.\n"
#     "10. Enjoy using the AI Web Scraper!"
#             "")

# st.markdown("### Note: "
#     "This application uses AI to parse content from websites. "
#     "Ensure you have the necessary permissions to scrape the website content. "
#     "The AI model may not always return accurate results, so review the output carefully. "
#     "If you have any questions or feedback, please reach out to the developer."
#     "Thank you for using the AI Web Scraper!")

# links = """
# <a href='https://www.google.com' target='_blank'>Google</a><br>
# <a href='https://www.github.com' target='_blank'>GitHub</a><br>
# <a href='https://www.python.org' target='_blank'>Python</a>
# """

# st.markdown(links, unsafe_allow_html=True);




# st.markdown("""
# <style>
#     .stButton button {
#         background-color: #4CAF50; /* Green */
#         border: none;
#         color: white;
#         padding: 10px 20px;
#         text-align: center;
#         text-decoration: none;
#         display: inline-block;
#         font-size: 16px;
#         margin: 4px 2px;
#         cursor: pointer;
#     }   
            
# """)